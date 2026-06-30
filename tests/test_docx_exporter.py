"""
Word 导出模块测试。

策略：
- 生成字节流后用 python-docx 重新解析，验证内容正确性
- 验证：标题层级、正文保留、表格、列表、空输入不崩溃
"""
import io
import pytest

try:
    from docx import Document
    from src.docx_exporter import markdown_to_docx_bytes
    _AVAILABLE = True
except ImportError:
    _AVAILABLE = False

pytestmark = pytest.mark.skipif(not _AVAILABLE, reason="python-docx 未安装")


def parse_docx(data: bytes) -> Document:
    return Document(io.BytesIO(data))


def all_text(doc: Document) -> str:
    return " ".join(p.text for p in doc.paragraphs)


# ==========================================
# 基本输出
# ==========================================

class TestBasicOutput:

    def test_returns_bytes(self):
        result = markdown_to_docx_bytes("# 标题\n\n正文")
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_output_is_valid_docx(self):
        data = markdown_to_docx_bytes("# 标题\n\n正文")
        assert parse_docx(data) is not None

    def test_empty_input_does_not_crash(self):
        result = markdown_to_docx_bytes("")
        assert isinstance(result, bytes)

    def test_title_parameter_appears_in_doc(self):
        data = markdown_to_docx_bytes("正文内容", title="我的报告")
        doc = parse_docx(data)
        assert "我的报告" in all_text(doc)


# ==========================================
# 标题层级
# ==========================================

class TestHeadings:

    def _heading_texts(self, doc: Document) -> list:
        return [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]

    def test_h1_rendered_as_heading(self):
        doc = parse_docx(markdown_to_docx_bytes("# 一级标题"))
        assert any("一级标题" in t for t in self._heading_texts(doc))

    def test_h2_rendered_as_heading(self):
        doc = parse_docx(markdown_to_docx_bytes("## 二级标题"))
        assert any("二级标题" in t for t in self._heading_texts(doc))

    def test_h3_rendered_as_heading(self):
        doc = parse_docx(markdown_to_docx_bytes("### 三级标题"))
        assert any("三级标题" in t for t in self._heading_texts(doc))


# ==========================================
# 内容保留
# ==========================================

class TestContentPreservation:

    def test_paragraph_text_preserved(self):
        doc = parse_docx(markdown_to_docx_bytes("这是一段普通正文，应该被完整保留。"))
        assert "普通正文" in all_text(doc)

    def test_list_items_preserved(self):
        doc = parse_docx(markdown_to_docx_bytes("- 第一条\n- 第二条\n- 第三条"))
        text = all_text(doc)
        assert "第一条" in text
        assert "第二条" in text
        assert "第三条" in text

    def test_multiple_sections_all_preserved(self):
        md = "# 背景\n\n背景内容。\n\n## 方法\n\n方法内容。\n\n## 结论\n\n结论内容。"
        doc = parse_docx(markdown_to_docx_bytes(md))
        text = all_text(doc)
        assert "背景内容" in text
        assert "方法内容" in text
        assert "结论内容" in text

    def test_chinese_content_preserved(self):
        chinese = "人工智能技术正在快速发展，深度学习已成为核心方法。"
        doc = parse_docx(markdown_to_docx_bytes(chinese))
        assert "人工智能" in all_text(doc)


# ==========================================
# 表格
# ==========================================

class TestTable:

    def test_markdown_table_creates_word_table(self):
        md = "| 方法 | 准确率 |\n| --- | --- |\n| BERT | 92% |\n| GPT-4 | 96% |"
        doc = parse_docx(markdown_to_docx_bytes(md))
        assert len(doc.tables) >= 1

    def test_table_content_preserved(self):
        md = "| 方法 | 准确率 |\n| --- | --- |\n| BERT | 92% |"
        doc = parse_docx(markdown_to_docx_bytes(md))
        if doc.tables:
            table_text = " ".join(
                cell.text for row in doc.tables[0].rows for cell in row.cells
            )
            assert "BERT" in table_text
            assert "92%" in table_text


# ==========================================
# 长文档性能
# ==========================================

class TestLargeDocument:

    def test_large_document_does_not_crash(self):
        sections = []
        for i in range(10):
            sections.append(
                f"## 第{i+1}节\n\n{'这是本节的详细内容，包含分析与讨论。' * 15}"
            )
        data = markdown_to_docx_bytes("\n\n".join(sections), title="完整研究报告")
        assert len(data) > 1000

    def test_many_headings_no_crash(self):
        md = "\n\n".join(f"### 子节 {i}" for i in range(30))
        data = markdown_to_docx_bytes(md)
        assert isinstance(data, bytes)
