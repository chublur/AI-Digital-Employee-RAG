"""
Word 导出模块。

职责：
将 Markdown 格式的报告文本转为 .docx 文件的字节流，
供前端（Gradio gr.File / API 二进制响应）直接下载使用。

转换规则：
  # 标题    → Heading 1
  ## 标题   → Heading 2
  ### 标题  → Heading 3
  | 表格 |  → Word 表格
  **粗体**  → Bold Run
  - 列表项  → List Bullet 样式
  普通段落  → Normal 样式

为什么单独一个模块？
  docx 导出是可选功能，不安装 python-docx 时其他功能不受影响；
  单独模块也方便单独测试。
"""
import io
import logging
import re
from typing import List

logger = logging.getLogger(__name__)


def markdown_to_docx_bytes(markdown_text: str, title: str = "") -> bytes:
    """
    将 Markdown 字符串转换为 .docx 文件的字节流。

    Args:
        markdown_text: Markdown 格式的报告正文
        title:         文档标题（显示在正文最顶部，可为空）

    Returns:
        .docx 文件的 bytes，可直接写入文件或传给 st.download_button
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx 未安装，请运行: pip install python-docx>=1.1.0")

    doc = Document()

    # 文档标题
    if title:
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # ── 标题 ──────────────────────────────────────────
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)

        # ── Markdown 表格 ─────────────────────────────────
        elif line.startswith("|") and "|" in line[1:]:
            # 收集连续的表格行
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _add_table(doc, table_lines)
            continue  # 不执行末尾的 i += 1

        # ── 无序列表 ──────────────────────────────────────
        elif line.startswith("- ") or line.startswith("* "):
            text = line[2:].strip()
            para = doc.add_paragraph(style="List Bullet")
            _add_inline_formatting(para, text)

        # ── 引用块（> 开头）────────────────────────────────
        elif line.startswith("> "):
            para = doc.add_paragraph(line[2:].strip(), style="Quote")

        # ── 水平分隔线 ────────────────────────────────────
        elif line.strip() in ("---", "***", "___"):
            doc.add_paragraph("─" * 40)

        # ── 空行 ──────────────────────────────────────────
        elif not line.strip():
            pass  # 跳过空行，不额外加段落

        # ── 普通段落 ──────────────────────────────────────
        else:
            para = doc.add_paragraph()
            _add_inline_formatting(para, line.strip())

        i += 1

    # 写入内存缓冲区，返回字节
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _add_table(doc, table_lines: List[str]) -> None:
    """
    将 Markdown 表格行列表写入 Word 表格。
    自动跳过分隔行（| --- | --- |）。
    """

    # 解析行：过滤分隔行
    data_rows = []
    for line in table_lines:
        if re.match(r"^\|[\s\-:|]+\|$", line.strip()):
            continue  # 跳过 | --- | --- | 分隔行
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        data_rows.append(cells)

    if not data_rows:
        return

    # 对齐列数
    col_count = max(len(row) for row in data_rows)
    for row in data_rows:
        while len(row) < col_count:
            row.append("")

    try:
        table = doc.add_table(rows=len(data_rows), cols=col_count)
        table.style = "Table Grid"

        for r_idx, row_data in enumerate(data_rows):
            row = table.rows[r_idx]
            for c_idx, cell_text in enumerate(row_data):
                cell = row.cells[c_idx]
                cell.text = cell_text
                # 首行加粗（表头）
                if r_idx == 0:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
    except Exception as e:
        logger.warning(f"Word 表格写入失败，改为纯文本: {e}")
        for row in data_rows:
            doc.add_paragraph("  |  ".join(row))


def _add_inline_formatting(para, text: str) -> None:
    """
    解析行内 **粗体** 和 *斜体* 标记，添加对应 Run 格式。
    不支持嵌套，遇到复杂情况降级为纯文本。
    """
    # 简单状态机：分割粗体/斜体片段
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")
    last = 0
    for m in pattern.finditer(text):
        # 普通文本
        if m.start() > last:
            para.add_run(text[last:m.start()])

        matched = m.group(0)
        if matched.startswith("**"):
            run = para.add_run(m.group(2))
            run.bold = True
        elif matched.startswith("*"):
            run = para.add_run(m.group(3))
            run.italic = True
        elif matched.startswith("`"):
            run = para.add_run(m.group(4))
            run.font.name = "Courier New"

        last = m.end()

    # 剩余普通文本
    if last < len(text):
        para.add_run(text[last:])
